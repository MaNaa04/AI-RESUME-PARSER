"""
parser.py
---------
Handles text extraction from uploaded resume files.
Supports .pdf (via pdfplumber) and .docx (via python-docx).
"""

import io
import pdfplumber
import docx


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extract all text from a PDF file given its raw bytes.

    Args:
        file_bytes: Raw bytes of the PDF file.

    Returns:
        Extracted text as a single string.

    Raises:
        ValueError: If the PDF has no extractable text content.
        Exception: For other PDF-related parsing errors.
    """
    text_parts = []

    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            if len(pdf.pages) == 0:
                raise ValueError("The PDF file appears to be empty (no pages found).")

            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
    except ValueError:
        raise
    except Exception as e:
        raise Exception(f"Failed to parse PDF: {str(e)}")

    full_text = "\n".join(text_parts).strip()

    if not full_text:
        raise ValueError(
            "No readable text could be extracted from this PDF. "
            "It may be a scanned image-based PDF. Please upload a text-based PDF."
        )

    return full_text


def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extract all text from a DOCX file given its raw bytes.

    Args:
        file_bytes: Raw bytes of the DOCX file.

    Returns:
        Extracted text as a single string.

    Raises:
        ValueError: If the DOCX file is empty or unreadable.
        Exception: For other DOCX-related parsing errors.
    """
    try:
        document = docx.Document(io.BytesIO(file_bytes))
    except Exception as e:
        raise Exception(f"Failed to open DOCX file: {str(e)}")

    paragraphs = [para.text for para in document.paragraphs if para.text.strip()]

    # Also extract text from tables
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text and cell_text not in paragraphs:
                    paragraphs.append(cell_text)

    full_text = "\n".join(paragraphs).strip()

    if not full_text:
        raise ValueError("The DOCX file appears to be empty or contains no readable text.")

    return full_text
